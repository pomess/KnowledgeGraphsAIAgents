"""
File parsers for multiple document formats.
Each parser returns a list of LangChain Document objects.
"""
import csv
import io
from typing import List
from langchain_core.documents import Document


def parse_txt(content: bytes, source: str) -> List[Document]:
    """Parse plain text files."""
    text = content.decode("utf-8")
    return [Document(page_content=text, metadata={"source": source})]


def parse_pdf(content: bytes, source: str) -> List[Document]:
    """Parse PDF files using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF parsing. Install with: pip install pymupdf")

    documents = []
    pdf = fitz.open(stream=content, filetype="pdf")
    full_text = []
    for page in pdf:
        full_text.append(page.get_text())
    pdf.close()

    text = "\n".join(full_text).strip()
    if text:
        documents.append(Document(page_content=text, metadata={"source": source}))
    return documents


def parse_docx(content: bytes, source: str) -> List[Document]:
    """Parse DOCX files using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if text:
        return [Document(page_content=text, metadata={"source": source})]
    return []


def parse_csv(content: bytes, source: str) -> List[Document]:
    """Parse CSV files — concatenates all rows into a readable text format."""
    text_content = content.decode("utf-8")
    reader = csv.DictReader(io.StringIO(text_content))

    rows = []
    for row in reader:
        row_parts = [f"{key}: {value}" for key, value in row.items() if value]
        rows.append(". ".join(row_parts))

    text = "\n".join(rows)
    if text:
        return [Document(page_content=text, metadata={"source": source})]
    return []


def parse_file(content: bytes, filename: str) -> List[Document]:
    """Route to the correct parser based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "txt":
        return parse_txt(content, filename)
    elif ext == "pdf":
        return parse_pdf(content, filename)
    elif ext in ("docx", "doc"):
        return parse_docx(content, filename)
    elif ext == "csv":
        return parse_csv(content, filename)
    else:
        # Fallback: try to decode as UTF-8 text
        try:
            text = content.decode("utf-8")
            return [Document(page_content=text, metadata={"source": filename})]
        except Exception:
            raise ValueError(f"Unsupported file format: .{ext}")
